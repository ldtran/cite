import streamlit as st
import requests
from bs4 import BeautifulSoup
import re
from io import BytesIO
import pdfplumber
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

# Page config
st.set_page_config(
    page_title="Citation Generator",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom styling
st.markdown("""
<style>
    .main { padding: 2rem; }
    .stTextArea textarea { font-family: monospace; }
</style>
""", unsafe_allow_html=True)

# ================================================================
# URL DETECTION FUNCTION
# ================================================================

def detect_citation_type(url):
    """
    Automatically detect if URL is an AG Annual Report, Ontario News Release, or Ontario e-Laws.
    Returns 'ag', 'news', 'elaws', or None if unrecognized.
    """
    if 'auditor.on.ca' in url:
        return 'ag'
    elif 'news.ontario.ca' in url:
        return 'news'
    elif 'ontario.ca/laws/statute' in url or 'ontario.ca/laws/regulation' in url:
        return 'elaws'
    else:
        return None

# ================================================================
# AG ANNUAL REPORT FUNCTIONS
# ================================================================

def is_pdf_url(url):
    return url.lower().endswith('.pdf')

def extract_year_from_url(url):
    match = re.search(r'(?:en|ar|fr)(\d{2,4})', url)
    if match:
        year = match.group(1)
        if len(year) == 2:
            year_int = int(year)
            if year_int >= 50:
                year = '19' + year
            else:
                year = '20' + year
        return year
    
    match = re.search(r'19\d{2}|20\d{2}', url)
    if match:
        return match.group(0)
    
    return None

def extract_chapter_section_from_filename(filename):
    match = re.search(r'v\d+_(\d)(\d{2})', filename)
    if match:
        chapter = match.group(1)
        section_digits = match.group(2)
        if section_digits != '00':
            return chapter, f"{chapter}.{section_digits}"
        else:
            return chapter, None
    
    match = re.search(r'^(\d)(\d{2})en\d{2}', filename)
    if match:
        chapter = match.group(1)
        section_digits = match.group(2)
        if section_digits != '00':
            return chapter, f"{chapter}.{section_digits}"
        else:
            return chapter, None
    
    return None, None

def extract_title_from_pdf_metadata(url):
    try:
        headers = {'User-Agent': 'Mozilla/5.0'}
        response = requests.get(url, headers=headers, timeout=30)
        response.raise_for_status()
        
        pdf_file = BytesIO(response.content)
        with pdfplumber.open(pdf_file) as pdf:
            metadata = pdf.metadata
            if metadata and 'Title' in metadata and metadata['Title']:
                title = metadata['Title']
                if isinstance(title, bytes):
                    title = title.decode('utf-8', errors='ignore')
                title = title.strip()
                title = re.sub(r'^\d{4}\s+(?:Provincial\s+)?Auditor\'?s?\s+Report:\s*', '', title, flags=re.I)
                title = re.sub(r'^VFM\s+\d+\.\d{2}\s*:\s*', '', title, flags=re.I)
                title = re.sub(r'^\d+\.\d{2}\s*:\s*', '', title)
                title = re.sub(r'^\d+\.\d{2}\s+', '', title)
                return title if title else None
            return None
    except:
        return None

def extract_title_from_html(url):
    try:
        headers = {'User-Agent': 'Mozilla/5.0'}
        response = requests.get(url, headers=headers, timeout=10)
        response.raise_for_status()
        
        soup = BeautifulSoup(response.content, 'html.parser')
        for tag in ['h1', 'h2', 'h3']:
            heading = soup.find(tag)
            if heading and heading.text.strip():
                title = heading.text.strip()
                if "Office of the Auditor General" not in title and "Auditor General" not in title:
                    return title
        
        content_area = soup.find(['div', 'section'], class_=re.compile('content|main|article', re.I))
        if content_area:
            heading = content_area.find(['h1', 'h2', 'h3'])
            if heading and heading.text.strip():
                title = heading.text.strip()
                if "Office of the Auditor General" not in title:
                    return title
        
        return None
    except:
        return None

def fetch_ag_citation(url):
    year = extract_year_from_url(url)
    if not year:
        return None, "Could not extract year"
    
    org_name = "Office of the Auditor General of Ontario"
    year_report = f"{year} Annual Report"
    
    if is_pdf_url(url):
        title = extract_title_from_pdf_metadata(url)
        if not title:
            return None, "Could not extract title from PDF"
        
        filename = url.split('/')[-1]
        chapter, section = extract_chapter_section_from_filename(filename)
        title = ' '.join(title.split())
        
        if chapter and section:
            citation = f'{org_name}, "[{title}]({url})", Chapter {chapter}, Section {section}, *{year_report}*.'
        elif chapter:
            citation = f'{org_name}, "[{title}]({url})", Chapter {chapter}, *{year_report}*.'
        else:
            citation = f'{org_name}, "[{title}]({url})", *{year_report}*.'
    else:
        title = extract_title_from_html(url)
        if not title:
            return None, "Could not extract title from page"
        
        title = ' '.join(title.split())
        citation = f'{org_name}, "[{title}]({url})", *{year_report}*.'
    
    return citation, None

# ================================================================
# ONTARIO NEWS RELEASE FUNCTIONS
# ================================================================

def extract_release_id(url):
    """Extract the numeric release ID from an Ontario news release URL."""
    parts = url.split('/')
    for part in parts:
        if part.isdigit():
            return part
    return None

def format_ministries(main_ministry, partner_ministries):
    """Format multiple ministries with proper comma placement and final 'and'."""
    ministries = [main_ministry] + [m['name'] for m in partner_ministries if 'name' in m and m['name'] != main_ministry]
    ministries = list(dict.fromkeys(ministries))
    
    if len(ministries) == 1:
        return ministries[0]
    elif len(ministries) == 2:
        return f"{ministries[0]} and {ministries[1]}"
    else:
        return ", ".join(ministries[:-1]) + f", and {ministries[-1]}"

def fetch_news_release_citation(url):
    """Fetch data from Ontario API and generate formatted citation."""
    release_id = extract_release_id(url)
    if not release_id:
        return None, f"Could not extract release ID from URL"
    
    api_url = f"https://api.news.ontario.ca/api/v1/releases/{release_id}?language=en"
    
    try:
        response = requests.get(api_url, timeout=10)
        response.raise_for_status()
    except requests.RequestException as e:
        return None, f"Error fetching data: {e}"
    
    try:
        data = response.json()['data']
        main_ministry = data.get('ministry_name', '')
        partner_ministries = data.get('partner_ministries', [])
        ministries_str = format_ministries(main_ministry, partner_ministries)
        title = data.get('content_title', 'No title')
        release_type = data.get('release_type_name', 'News Release').lower()
        date_str = data.get('release_date_time_formatted', '')
        
        citation = f'{ministries_str}, "[{title}]({url})", *{release_type}*, {date_str}.'
        return citation, None
    except Exception as e:
        return None, f"Error parsing data: {e}"

# ================================================================
# ONTARIO E-LAWS FUNCTIONS (STATUTES & REGULATIONS)
# ================================================================

def fetch_elaws_citation(url):
    """
    Extract titles from Ontario e-Laws (statutes or regulations),
    and extract section reference if present in URL anchor/parameters.
    
    For statutes: Returns both full and short titles
    For regulations: Returns only short title (O. Reg. X/YY)
    """
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        }
        response = requests.get(url, headers=headers, timeout=10)
        response.raise_for_status()
        
        soup = BeautifulSoup(response.content, 'html.parser')
        
        # Extract the full title from page title
        page_title = soup.title.string if soup.title else ""
        
        # Determine if this is a statute or regulation
        is_regulation = 'regulation' in url
        
        if page_title:
            # Remove trailing descriptions after the main legislation title
            full_title_match = re.search(r'^([^-]+)', page_title)
            full_title = full_title_match.group(1).strip() if full_title_match else page_title
            
            if is_regulation:
                # For regulations, extract ONLY short form: "O. Reg. 221/08"
                short_match = re.match(r'^(O\.\s*Reg\.\s*\d+/\d+)', full_title)
                short_title = short_match.group(1) if short_match else full_title
                # For regs, we only return short title
                full_title = None
            else:
                # For statutes, extract short title: "Early Childhood Educators Act, 2007"
                short_title_match = re.match(r'^([^,]+, \d{4})', full_title)
                short_title = short_title_match.group(1) if short_title_match else full_title
        else:
            full_title = "Legislation"
            short_title = "Legislation"
        
        # Extract section from URL anchor (#)
        section = None
        
        if '#' in url:
            anchor = url.split('#')[1].split('?')[0]
            if anchor and anchor.startswith('BK'):
                anchor_element = soup.find(attrs={'id': anchor})
                if anchor_element:
                    for parent in anchor_element.parents:
                        section_heading = parent.find(['h2', 'h3', 'strong'])
                        if section_heading:
                            section_text = section_heading.get_text().strip()
                            section_match = re.match(r'^(\d+)', section_text)
                            if section_match:
                                section = section_match.group(1)
                                break
        
        # Remove query parameters and anchors from URL for clean citation
        clean_url = url.split('?')[0].split('#')[0]
        
        return {
            'full_title': full_title,
            'short_title': short_title,
            'section': section,
            'clean_url': clean_url,
            'is_regulation': is_regulation,
            'full_citation': f"[{full_title}]({clean_url})" + (f", s. {section}." if section else ".") if full_title else None,
            'short_citation': f"[{short_title}]({clean_url})" + (f", s. {section}." if section else ".")
        }
    
    except Exception as e:
        return {
            'error': str(e),
            'full_title': None,
            'short_title': None,
            'section': None,
            'is_regulation': 'regulation' in url
        }

# ================================================================
# SHARED DOCX FUNCTIONS
# ================================================================

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, reltype="http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)
    
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Arial')
    rFonts.set(qn('w:hAnsi'), 'Arial')
    rPr.append(rFonts)
    
    new_run.append(rPr)
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)
    
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    
    return hyperlink

def generate_docx(citation_texts):
    doc = Document()
    
    for citation_md in citation_texts:
        p = doc.add_paragraph()
        pattern = re.compile(r'(\[.*?\]\(.*?\)|\*.*?\*|_.*?_)')
        parts = pattern.split(citation_md)
        
        for part in parts:
            if not part:
                continue
            
            if part.startswith('[') and '](' in part and part.endswith(')'):
                link_text = re.findall(r'\[(.*?)\]', part)[0]
                link_url = re.findall(r'\((.*?)\)', part)[0]
                add_hyperlink(p, link_text, link_url)
            elif (part.startswith('*') and part.endswith('*')) or (part.startswith('_') and part.endswith('_')):
                run = p.add_run(part[1:-1])
                run.italic = True
                run.font.name = 'Arial'
                run.font.size = Pt(11)
            else:
                run = p.add_run(part)
                run.font.name = 'Arial'
                run.font.size = Pt(11)
        
        p.add_run('\n')
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ================================================================
# STREAMLIT UI
# ================================================================

st.title("📄 Citation Generator")
st.markdown("Generate formatted citations for a variety of sources automatically.")

with st.sidebar:
    st.header("⚙️ Settings")
    st.divider()
    
    st.subheader("💡 Instructions")
    st.markdown("""
1. Paste URLs (AG Annual Reports, Ontario News Releases, or e-Laws)
2. Click "Generate Citations"
3. Download as DOCX

**Supported:**
- AG Annual Reports (HTML, PDF)
- Ontario News Releases (HTML)
- Ontario e-Laws Statutes & Regulations
- Mix any types in one batch
    """)
    
    st.divider()
    
    st.subheader("📚 Example URLs")
    with st.expander("View examples"):
        st.code("""https://www.auditor.on.ca/en/content/annualreports/audits/en2024/AR-PA_ONlandtribunal_en24.html

https://news.ontario.ca/en/release/1006488/ontario-welcomes-207-million-investment-in-the-advanced-manufacturing-sector

https://www.auditor.on.ca/en/content/annualreports/arreports/en18/v1_306en18.pdf

https://www.ontario.ca/laws/statute/07e07

https://www.ontario.ca/laws/regulation/080221""")

st.markdown("### 📋 Paste Your URLs")

urls_input = st.text_area(
    "Enter URLs (AG Annual Reports, Ontario News Releases, or e-Laws - one per line):",
    height=200,
    placeholder="Paste URLs here...",
    label_visibility="collapsed"
)

col1, col2 = st.columns([1, 3])

with col1:
    generate_btn = st.button("🚀 Generate Citations", use_container_width=True)

if generate_btn:
    if not urls_input.strip():
        st.warning("⚠️ Please enter at least one URL")
    else:
        urls = [url.strip() for url in urls_input.split('\n') if url.strip()]
        
        st.markdown("---")
        st.markdown("### 📝 Results")
        
        citations = []
        progress_bar = st.progress(0)
        status_container = st.container()
        
        for i, url in enumerate(urls):
            progress = (i + 1) / len(urls)
            progress_bar.progress(progress)
            
            with status_container:
                with st.spinner(f"Processing [{i+1}/{len(urls)}]: {url[:60]}..."):
                    # Auto-detect citation type
                    citation_type = detect_citation_type(url)
                    
                    if citation_type == 'ag':
                        citation, error = fetch_ag_citation(url)
                        if error:
                            st.error(f"❌ {url}\n{error}")
                        else:
                            citations.append((url, citation, 'ag'))
                            st.success(f"✓ {url[:70]}...")
                    
                    elif citation_type == 'news':
                        citation, error = fetch_news_release_citation(url)
                        if error:
                            st.error(f"❌ {url}\n{error}")
                        else:
                            citations.append((url, citation, 'news'))
                            st.success(f"✓ {url[:70]}...")
                    
                    elif citation_type == 'elaws':
                        citation_data = fetch_elaws_citation(url)
                        if 'error' in citation_data:
                            st.error(f"❌ {url}\n{citation_data['error']}")
                        else:
                            # Display results for e-Laws
                            st.success(f"✓ {url[:70]}...")
                            
                            # Show both full and short for statutes, only short for regs
                            if citation_data['full_title']:
                                st.markdown("**Full Title Citation:**")
                                st.code(citation_data['full_citation'], language="markdown")
                                st.markdown("**Short Title Citation:**")
                                st.code(citation_data['short_citation'], language="markdown")
                                citations.append((url, citation_data['full_citation'], 'elaws'))
                            else:
                                st.markdown("**Citation:**")
                                st.code(citation_data['short_citation'], language="markdown")
                                citations.append((url, citation_data['short_citation'], 'elaws'))
                            
                            if citation_data['section']:
                                st.caption(f"📌 Section Referenced: s. {citation_data['section']}")
                    
                    else:
                        st.error(f"❌ {url}\nURL is not recognized. Please check the URL format.")
        
        progress_bar.empty()
        
        if citations:
            all_citation_texts = [c[1] for c in citations]
            docx_buffer = generate_docx(all_citation_texts)
            
            st.download_button(
                label="📥 Download All Citations (Word DOCX)",
                data=docx_buffer,
                file_name="Citations.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
